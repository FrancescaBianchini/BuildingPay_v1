# -*- coding: utf-8 -*-
import base64
import logging
from io import BytesIO
from odoo import http, _
from odoo.http import request
from odoo.addons.portal.controllers.portal import CustomerPortal, pager as portal_pager
from odoo.exceptions import AccessError, UserError

_logger = logging.getLogger(__name__)


class BuildingPayPortal(CustomerPortal):
    """
    Controller portale BuildingPay.
    Aggiunge sezioni:
    - Contratto Generale: download/upload Accordo Condomini Aggregati
    - Condomini: CRUD indirizzi di tipo 'condominio'
    """

    # -------------------------------------------------------
    # Home portale: aggiunge le sezioni BuildingPay
    # -------------------------------------------------------
    def _prepare_home_portal_values(self, counters):
        values = super()._prepare_home_portal_values(counters)
        partner = request.env.user.partner_id

        if 'condominio_count' in counters:
            values['condominio_count'] = request.env['res.partner'].sudo().search_count([
                ('parent_id', '=', partner.id),
                ('type', '=', 'condominio'),
                ('active', '=', True),
            ])

        return values

    # -------------------------------------------------------
    # SEZIONE: Contratto Generale
    # -------------------------------------------------------
    @http.route('/my/contratto', type='http', auth='user', website=True)
    def portal_contratto_generale(self, **kw):
        """Pagina del contratto generale nel portale."""
        partner = request.env.user.partner_id
        config = request.env['buildingpay.config'].sudo().get_config_for_website()

        values = {
            'partner': partner,
            'config': config,
            'page_name': 'contratto',
        }
        return request.render('buildingpay.portal_contratto', values)

    @http.route('/my/contratto/download', type='http', auth='user', website=True)
    def portal_contratto_download(self, **kw):
        """
        Download del template 'Accordo Condomini Aggregati' con sostituzione
        dei placeholder:
        - [NOME AMMINISTRATORE] → nome e cognome utente portale
        - [________] → codice fiscale utente portale
        """
        partner = request.env.user.partner_id
        config = request.env['buildingpay.config'].sudo().get_config_for_website()

        if not config or not config.contratto_template:
            return request.not_found()

        try:
            from docx import Document

            # Carica il template dalla configurazione
            template_data = base64.b64decode(config.contratto_template)
            doc = Document(BytesIO(template_data))

            nome_amministratore = partner.name or ''
            codice_fiscale = partner.fiscalcode or ''

            # Sostituzione nei paragrafi
            for paragraph in doc.paragraphs:
                self._replace_placeholder_in_paragraph(
                    paragraph, '[NOME AMMINISTRATORE]', nome_amministratore)
                self._replace_placeholder_in_paragraph(
                    paragraph, '[________]', codice_fiscale)

            # Sostituzione nelle tabelle
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            self._replace_placeholder_in_paragraph(
                                paragraph, '[NOME AMMINISTRATORE]', nome_amministratore)
                            self._replace_placeholder_in_paragraph(
                                paragraph, '[________]', codice_fiscale)

            # Genera il file in memoria
            output = BytesIO()
            doc.save(output)
            output.seek(0)
            file_data = output.read()

            filename = 'Accordo Condomini Aggregati.docx'
            return request.make_response(
                file_data,
                headers=[
                    ('Content-Type',
                     'application/vnd.openxmlformats-officedocument.wordprocessingml.document'),
                    ('Content-Disposition',
                     'attachment; filename="%s"' % filename),
                    ('Content-Length', len(file_data)),
                ],
            )

        except Exception as e:
            _logger.error('BuildingPay: errore download contratto: %s', e)
            return request.redirect('/my/contratto?error=download_error')

    def _replace_placeholder_in_paragraph(self, paragraph, placeholder, replacement):
        """
        Sostituisce un placeholder nel testo di un paragrafo docx,
        preservando la formattazione dei run.
        """
        if placeholder not in paragraph.text:
            return

        # Concatena tutti i testi per trovare il placeholder anche se
        # distribuito su più run
        full_text = ''.join(run.text for run in paragraph.runs)
        if placeholder not in full_text:
            return

        new_text = full_text.replace(placeholder, replacement)

        # Riscrivi: metti tutto il testo nel primo run e svuota gli altri
        if paragraph.runs:
            paragraph.runs[0].text = new_text
            for run in paragraph.runs[1:]:
                run.text = ''

    @http.route('/my/contratto/upload', type='http', auth='user', website=True,
                methods=['POST'])
    def portal_contratto_upload(self, **kw):
        """
        Upload del file contratto firmato dall'utente portale.
        Attiva il flag accordo_condomini_aggregati.
        """
        partner = request.env.user.partner_id
        uploaded_file = kw.get('contratto_file')

        if not uploaded_file:
            return request.redirect('/my/contratto?error=no_file')

        try:
            file_data = uploaded_file.read()
            filename = uploaded_file.filename

            file_b64 = base64.b64encode(file_data)
            partner.sudo().action_upload_contratto(file_b64, filename)

            return request.redirect('/my/contratto?success=1')
        except Exception as e:
            _logger.error('BuildingPay: errore upload contratto: %s', e)
            return request.redirect('/my/contratto?error=upload_error')

    # -------------------------------------------------------
    # SEZIONE: Condomini
    # -------------------------------------------------------
    @http.route('/my/condomini', type='http', auth='user', website=True)
    def portal_condomini_list(self, page=1, **kw):
        """Lista dei condomini dell'amministratore."""
        partner = request.env.user.partner_id

        if not partner.is_amministratore:
            return request.redirect('/my')

        domain = [
            ('parent_id', '=', partner.id),
            ('type', '=', 'condominio'),
            ('active', '=', True),
        ]

        condomini = request.env['res.partner'].sudo().search(domain)

        values = {
            'partner': partner,
            'condomini': condomini,
            'page_name': 'condomini',
        }
        return request.render('buildingpay.portal_condomini', values)

    @http.route('/my/condomini/new', type='http', auth='user', website=True)
    def portal_condominio_new(self, **kw):
        """Form per aggiungere un nuovo condominio."""
        partner = request.env.user.partner_id

        if not partner.is_amministratore:
            return request.redirect('/my')

        countries = request.env['res.country'].sudo().search([])
        values = {
            'partner': partner,
            'condominio': None,
            'countries': countries,
            'page_name': 'condomini_new',
            'mode': 'create',
        }
        return request.render('buildingpay.portal_condominio_form', values)

    @http.route('/my/condomini/new/save', type='http', auth='user', website=True,
                methods=['POST'])
    def portal_condominio_create(self, **kw):
        """Salva un nuovo indirizzo condominio."""
        partner = request.env.user.partner_id

        if not partner.is_amministratore:
            return request.redirect('/my')

        params = request.params
        errors = self._validate_condominio_form(params)

        if errors:
            countries = request.env['res.country'].sudo().search([])
            return request.render('buildingpay.portal_condominio_form', {
                'partner': partner,
                'condominio': None,
                'countries': countries,
                'errors': errors,
                'form_data': params,
                'mode': 'create',
                'page_name': 'condomini_new',
            })

        try:
            condominio_vals = self._prepare_condominio_vals(params, partner)
            condominio = request.env['res.partner'].sudo().create(condominio_vals)

            # Salva IBAN nei conti bancari
            iban = params.get('iban', '').strip()
            if iban:
                request.env['res.partner.bank'].sudo().create({
                    'partner_id': condominio.id,
                    'acc_number': iban,
                })

            # Attiva flag electronic invoice se codice destinatario presente
            if params.get('codice_destinatario'):
                condominio.sudo().write({
                    'electronic_invoice_subjected': True,
                    'electronic_invoice_obliged_subject': True,
                })

            return request.redirect('/my/condomini?success_add=1')
        except Exception as e:
            _logger.error('BuildingPay: errore creazione condominio: %s', e)
            return request.redirect('/my/condomini?error=create_error')

    @http.route('/my/condomini/<int:condominio_id>', type='http', auth='user', website=True)
    def portal_condominio_detail(self, condominio_id, **kw):
        """Dettaglio/modifica di un condominio esistente."""
        partner = request.env.user.partner_id
        condominio = self._get_condominio_or_redirect(condominio_id, partner)
        if isinstance(condominio, type(request.redirect('/'))):
            return condominio

        # Recupera IBAN dal conto bancario
        bank = request.env['res.partner.bank'].sudo().search([
            ('partner_id', '=', condominio.id),
        ], limit=1)

        countries = request.env['res.country'].sudo().search([])
        values = {
            'partner': partner,
            'condominio': condominio,
            'bank': bank,
            'countries': countries,
            'page_name': 'condomini_edit',
            'mode': 'edit',
        }
        return request.render('buildingpay.portal_condominio_form', values)

    @http.route('/my/condomini/<int:condominio_id>/save', type='http', auth='user',
                website=True, methods=['POST'])
    def portal_condominio_update(self, condominio_id, **kw):
        """Aggiorna un condominio esistente."""
        partner = request.env.user.partner_id
        condominio = self._get_condominio_or_redirect(condominio_id, partner)
        if isinstance(condominio, type(request.redirect('/'))):
            return condominio

        params = request.params
        errors = self._validate_condominio_form(params)

        if errors:
            countries = request.env['res.country'].sudo().search([])
            bank = request.env['res.partner.bank'].sudo().search([
                ('partner_id', '=', condominio.id),
            ], limit=1)
            return request.render('buildingpay.portal_condominio_form', {
                'partner': partner,
                'condominio': condominio,
                'bank': bank,
                'countries': countries,
                'errors': errors,
                'form_data': params,
                'mode': 'edit',
                'page_name': 'condomini_edit',
            })

        try:
            condominio_vals = self._prepare_condominio_vals(params, partner)
            # Non sovrascriviamo parent_id e type in aggiornamento
            condominio_vals.pop('parent_id', None)
            condominio_vals.pop('type', None)
            condominio.sudo().write(condominio_vals)

            # Aggiorna IBAN
            iban = params.get('iban', '').strip()
            existing_bank = request.env['res.partner.bank'].sudo().search([
                ('partner_id', '=', condominio.id),
            ], limit=1)
            if iban:
                if existing_bank:
                    existing_bank.sudo().write({'acc_number': iban})
                else:
                    request.env['res.partner.bank'].sudo().create({
                        'partner_id': condominio.id,
                        'acc_number': iban,
                    })

            # Aggiorna flag electronic invoice
            if params.get('codice_destinatario'):
                condominio.sudo().write({
                    'electronic_invoice_subjected': True,
                    'electronic_invoice_obliged_subject': True,
                })

            return request.redirect('/my/condomini?success_edit=1')
        except Exception as e:
            _logger.error('BuildingPay: errore aggiornamento condominio: %s', e)
            return request.redirect('/my/condomini?error=update_error')

    @http.route('/my/condomini/<int:condominio_id>/archive', type='http', auth='user',
                website=True, methods=['POST'])
    def portal_condominio_archive(self, condominio_id, **kw):
        """Archivia un condominio (lo rende non attivo)."""
        partner = request.env.user.partner_id
        condominio = self._get_condominio_or_redirect(condominio_id, partner)
        if isinstance(condominio, type(request.redirect('/'))):
            return condominio

        try:
            condominio.sudo().action_archive_condominio()
            return request.redirect('/my/condomini?success_archive=1')
        except Exception as e:
            _logger.error('BuildingPay: errore archiviazione condominio: %s', e)
            return request.redirect('/my/condomini?error=archive_error')

    # -------------------------------------------------------
    # Metodi di utilità
    # -------------------------------------------------------
    def _get_condominio_or_redirect(self, condominio_id, partner):
        """
        Verifica che il condominio esista e appartenga all'utente corrente.
        Ritorna il record condominio oppure un redirect.
        """
        condominio = request.env['res.partner'].sudo().browse(condominio_id)
        if (not condominio.exists() or
                condominio.parent_id.id != partner.id or
                condominio.type != 'condominio'):
            return request.redirect('/my/condomini')
        return condominio

    def _validate_condominio_form(self, params):
        """Valida i dati del form condominio. Ritorna dict degli errori."""
        errors = {}
        if not params.get('name', '').strip():
            errors['name'] = _('Il nome è obbligatorio.')
        if not params.get('street', '').strip():
            errors['street'] = _('L\'indirizzo è obbligatorio.')
        if not params.get('city', '').strip():
            errors['city'] = _('La città è obbligatoria.')
        if not params.get('zip', '').strip():
            errors['zip'] = _('Il CAP è obbligatorio.')
        if not params.get('fiscalcode', '').strip():
            errors['fiscalcode'] = _('Il codice fiscale è obbligatorio.')
        return errors

    def _prepare_condominio_vals(self, params, parent_partner):
        """Prepara il dict dei valori per creare/aggiornare un condominio."""
        vals = {
            'name': params.get('name', '').strip(),
            'type': 'condominio',
            'parent_id': parent_partner.id,
            'street': params.get('street', '').strip(),
            'street2': params.get('street2', '').strip(),
            'city': params.get('city', '').strip(),
            'zip': params.get('zip', '').strip(),
            'fiscalcode': params.get('fiscalcode', '').strip(),
            'pec_mail': params.get('pec_mail', '').strip(),
            'codice_destinatario': params.get('codice_destinatario', '').strip(),
        }
        country_id = params.get('country_id')
        if country_id:
            vals['country_id'] = int(country_id)
        state_id = params.get('state_id')
        if state_id:
            vals['state_id'] = int(state_id)
        return vals
